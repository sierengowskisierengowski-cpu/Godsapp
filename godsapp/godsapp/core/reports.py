"""Report generator: Markdown, HTML, JSON, SARIF (always available),
plus DOCX / XLSX / PDF when optional libraries are installed.

Author: Joseph Sierengowski.
"""
from __future__ import annotations

import html
import json
from datetime import datetime
from pathlib import Path
from typing import Any

from sqlalchemy import select
from sqlalchemy.orm import selectinload

from godsapp.core import paths
from godsapp.core.settings import load_settings
from godsapp.db import Evidence, Finding, Scan, Workspace, get_session


SUPPORTED_FORMATS = ["markdown", "html", "json", "sarif", "docx", "xlsx", "pdf"]


def default_output_dir() -> Path:
    s = load_settings()
    if s.reports.output_dir:
        return Path(s.reports.output_dir).expanduser()
    return paths.DATA_DIR / "reports"


def generate(workspace_id: str, fmt: str, out_path: Path | None = None) -> Path:
    fmt = (fmt or "markdown").lower()
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"unsupported format: {fmt}")
    snap = _snapshot(workspace_id)
    if out_path is None:
        out_dir = default_output_dir()
        out_dir.mkdir(parents=True, exist_ok=True)
        stamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
        ext = {"markdown": "md"}.get(fmt, fmt)
        slug = "".join(c if c.isalnum() else "-" for c in snap["workspace"]["name"]).strip("-")
        out_path = out_dir / f"{slug}-{stamp}.{ext}"
    else:
        out_path = Path(out_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

    if fmt == "json":
        out_path.write_text(json.dumps(snap, indent=2, default=str))
    elif fmt == "markdown":
        out_path.write_text(_render_md(snap))
    elif fmt == "html":
        out_path.write_text(_render_html(snap))
    elif fmt == "sarif":
        out_path.write_text(json.dumps(_render_sarif(snap), indent=2))
    elif fmt == "docx":
        _render_docx(snap, out_path)
    elif fmt == "xlsx":
        _render_xlsx(snap, out_path)
    elif fmt == "pdf":
        _render_pdf(snap, out_path)
    return out_path


# ── snapshot ──────────────────────────────────────────────────────────────
def _snapshot(workspace_id: str) -> dict[str, Any]:
    with get_session() as s:
        ws = s.get(Workspace, workspace_id)
        if not ws:
            raise ValueError(f"workspace not found: {workspace_id}")
        scans = s.execute(
            select(Scan).where(Scan.workspace_id == workspace_id)
            .options(selectinload(Scan.findings))
            .order_by(Scan.created_at.desc())
        ).scalars().all()
        evidence = s.execute(
            select(Evidence).where(Evidence.workspace_id == workspace_id)
            .order_by(Evidence.created_at.desc())
        ).scalars().all()
        settings = load_settings()
        return {
            "workspace": {
                "id": ws.id, "name": ws.name,
                "description": ws.description, "target": ws.target,
                "created_at": ws.created_at.isoformat() if ws.created_at else None,
            },
            "report": {
                "generated_at": datetime.utcnow().isoformat() + "Z",
                "author": settings.reports.author,
                "org": settings.reports.org,
                "watermark": settings.reports.watermark,
            },
            "totals": {
                "scans": len(scans),
                "findings": sum(len(sc.findings) for sc in scans),
                "evidence": len(evidence),
            },
            "scans": [
                {
                    "id": sc.id, "tool": sc.tool, "category": sc.category, "target": sc.target,
                    "status": sc.status, "exit_code": sc.exit_code,
                    "started_at": sc.started_at.isoformat() if sc.started_at else None,
                    "finished_at": sc.finished_at.isoformat() if sc.finished_at else None,
                    "findings": [
                        {
                            "id": f.id, "title": f.title, "severity": f.severity,
                            "host": f.host, "port": f.port, "service": f.service,
                            "protocol": f.protocol, "description": f.description,
                            "status": getattr(f, "status", None) or "open",
                            "cvss_score": getattr(f, "cvss_score", None),
                            "cve_ids": getattr(f, "cve_ids", None),
                            "mitre_technique": getattr(f, "mitre_technique", None),
                            "tags": getattr(f, "tags", None),
                            "data": f.data,
                        }
                        for f in sc.findings
                    ],
                }
                for sc in scans
            ],
            "evidence": [
                {
                    "id": e.id, "filename": e.filename, "sha256": e.sha256,
                    "size_bytes": e.size_bytes, "mime_type": e.mime_type,
                    "created_at": e.created_at.isoformat() if e.created_at else None,
                    "note": e.note,
                }
                for e in evidence
            ],
        }


# ── renderers ─────────────────────────────────────────────────────────────
def _render_md(snap: dict) -> str:
    ws = snap["workspace"]; rep = snap["report"]; tot = snap["totals"]
    out = [
        f"# GodsApp Report — {ws['name']}",
        "",
        f"- **Target**: `{ws['target'] or 'n/a'}`",
        f"- **Description**: {ws['description'] or '_(none)_'}",
        f"- **Generated**: {rep['generated_at']}",
        f"- **Author**: {rep['author']}" + (f" · {rep['org']}" if rep['org'] else ""),
        "",
        "## Summary",
        f"- Scans: **{tot['scans']}**",
        f"- Findings: **{tot['findings']}**",
        f"- Evidence items: **{tot['evidence']}**",
        "",
        f"## Scans ({tot['scans']})",
        "",
    ]
    for sc in snap['scans']:
        out.append(f"### `{sc['tool']}` → `{sc['target']}`  ·  {sc['category']}")
        out.append(f"- Status: **{sc['status']}**  ·  exit {sc['exit_code']}")
        out.append(f"- Started: {sc['started_at'] or '–'}  ·  Finished: {sc['finished_at'] or '–'}")
        out.append(f"- Findings: **{len(sc['findings'])}**")
        if sc['findings']:
            out.append("")
            for f in sc['findings']:
                tag = f"[{f['severity'].upper()}]"
                extra = []
                if f.get('cvss_score'): extra.append(f"CVSS {f['cvss_score']}")
                if f.get('cve_ids'):    extra.append(f"CVE {f['cve_ids']}")
                if f.get('mitre_technique'): extra.append(f"MITRE {f['mitre_technique']}")
                tail = (" · " + " · ".join(extra)) if extra else ""
                out.append(f"  - **{tag}** {f['title']}{tail}")
        out.append("")
    out.append(f"## Evidence ({tot['evidence']})")
    out.append("")
    for ev in snap['evidence']:
        out.append(f"- `{ev['filename']}` — sha256:{ev['sha256'][:16]}…  ·  {ev['size_bytes']} bytes")
    if rep['watermark']:
        out.append("")
        out.append(f"---\n*{rep['watermark']}*")
    return "\n".join(out)


_HTML_CSS = """
body{font-family:system-ui,-apple-system,Segoe UI,sans-serif;background:#1a140d;color:#f5ecd9;padding:32px;max-width:960px;margin:auto;}
h1,h2,h3{color:#ebd7af;}
h1{border-bottom:2px solid #4a3a28;padding-bottom:8px;}
a{color:#d4b87a;}
.scan{border:1px solid #4a3a28;padding:16px;margin:12px 0;border-radius:8px;background:rgba(245,236,217,.03);}
.kv{color:#bca989;font-size:0.9em;}
.sev-critical{color:#e35d5d;font-weight:700;}
.sev-high{color:#e89640;font-weight:600;}
.sev-medium{color:#d6c43c;}
.sev-low{color:#7fb5c9;}
.sev-info{color:#9aa39a;}
table{width:100%;border-collapse:collapse;margin:12px 0;}
th,td{border-bottom:1px solid #4a3a28;padding:6px 10px;text-align:left;}
th{color:#ebd7af;}
footer{margin-top:48px;color:#7a6852;font-style:italic;text-align:center;}
"""


def _render_html(snap: dict) -> str:
    ws = snap["workspace"]; rep = snap["report"]; tot = snap["totals"]
    e = html.escape
    parts = [
        f"<!doctype html><html lang=en><head><meta charset='utf-8'>",
        f"<title>GodsApp Report — {e(ws['name'])}</title>",
        f"<style>{_HTML_CSS}</style></head><body>",
        f"<h1>GodsApp Report — {e(ws['name'])}</h1>",
        f"<p class='kv'><b>Target:</b> {e(ws['target'] or 'n/a')}<br>",
        f"<b>Description:</b> {e(ws['description'] or '—')}<br>",
        f"<b>Generated:</b> {e(rep['generated_at'])}<br>",
        f"<b>Author:</b> {e(rep['author'])}" + (f" · {e(rep['org'])}" if rep['org'] else "") + "</p>",
        "<h2>Summary</h2>",
        f"<p>Scans: <b>{tot['scans']}</b> · Findings: <b>{tot['findings']}</b> · Evidence: <b>{tot['evidence']}</b></p>",
        f"<h2>Scans ({tot['scans']})</h2>",
    ]
    for sc in snap['scans']:
        parts.append(f"<div class='scan'><h3>{e(sc['tool'])} → {e(sc['target'])}</h3>")
        parts.append(f"<p class='kv'>Category: {e(sc['category'])} · Status: <b>{e(sc['status'])}</b> · exit {sc['exit_code']}<br>")
        parts.append(f"Started: {e(sc['started_at'] or '–')} · Finished: {e(sc['finished_at'] or '–')}</p>")
        if sc['findings']:
            parts.append("<table><thead><tr><th>Severity</th><th>Title</th><th>Host</th><th>CVSS</th><th>CVE</th><th>MITRE</th></tr></thead><tbody>")
            for f in sc['findings']:
                parts.append(
                    f"<tr><td class='sev-{e(f['severity'])}'>{e(f['severity'])}</td>"
                    f"<td>{e(f['title'])}</td><td>{e(f['host'] or '')}</td>"
                    f"<td>{e(str(f.get('cvss_score') or ''))}</td>"
                    f"<td>{e(f.get('cve_ids') or '')}</td>"
                    f"<td>{e(f.get('mitre_technique') or '')}</td></tr>"
                )
            parts.append("</tbody></table>")
        else:
            parts.append("<p class='kv'>(no findings)</p>")
        parts.append("</div>")
    parts.append(f"<h2>Evidence ({tot['evidence']})</h2>")
    if snap['evidence']:
        parts.append("<table><thead><tr><th>File</th><th>SHA-256</th><th>Size</th></tr></thead><tbody>")
        for ev in snap['evidence']:
            parts.append(
                f"<tr><td>{e(ev['filename'])}</td>"
                f"<td><code>{e(ev['sha256'])}</code></td>"
                f"<td>{ev['size_bytes']}</td></tr>"
            )
        parts.append("</tbody></table>")
    if rep['watermark']:
        parts.append(f"<footer>{e(rep['watermark'])}</footer>")
    parts.append("</body></html>")
    return "\n".join(parts)


def _render_sarif(snap: dict) -> dict[str, Any]:
    runs = []
    for sc in snap['scans']:
        results = []
        for f in sc['findings']:
            level = {"critical": "error", "high": "error", "medium": "warning",
                     "low": "note", "info": "note"}.get(f['severity'], "note")
            results.append({
                "ruleId": sc['tool'],
                "level": level,
                "message": {"text": f['title']},
                "properties": {
                    "severity": f['severity'], "host": f['host'], "port": f['port'],
                    "service": f['service'], "cvss_score": f.get('cvss_score'),
                    "cve_ids": f.get('cve_ids'), "mitre_technique": f.get('mitre_technique'),
                },
            })
        runs.append({
            "tool": {"driver": {"name": sc['tool'], "informationUri":
                                "https://github.com/sierengowskisierengowski-cpu/Godsapp"}},
            "results": results,
        })
    return {
        "$schema": "https://raw.githubusercontent.com/oasis-tcs/sarif-spec/master/Schemata/sarif-schema-2.1.0.json",
        "version": "2.1.0",
        "runs": runs,
    }


def _render_docx(snap: dict, out_path: Path) -> None:
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise RuntimeError("DOCX export requires `python-docx` (pip install python-docx).") from e
    ws = snap["workspace"]; rep = snap["report"]; tot = snap["totals"]
    doc = Document()
    doc.add_heading(f"GodsApp Report — {ws['name']}", 0)
    doc.add_paragraph(f"Target: {ws['target'] or 'n/a'}")
    doc.add_paragraph(f"Generated: {rep['generated_at']}")
    doc.add_paragraph(f"Author: {rep['author']}" + (f" · {rep['org']}" if rep['org'] else ""))
    doc.add_heading("Summary", 1)
    doc.add_paragraph(f"Scans: {tot['scans']}  ·  Findings: {tot['findings']}  ·  Evidence: {tot['evidence']}")
    for sc in snap['scans']:
        doc.add_heading(f"{sc['tool']} → {sc['target']}", 2)
        doc.add_paragraph(f"Category: {sc['category']}  ·  Status: {sc['status']}  ·  exit {sc['exit_code']}")
        for f in sc['findings']:
            doc.add_paragraph(f"[{f['severity'].upper()}] {f['title']}", style="List Bullet")
    doc.save(str(out_path))


def _render_xlsx(snap: dict, out_path: Path) -> None:
    try:
        from openpyxl import Workbook
    except ImportError as e:
        raise RuntimeError("XLSX export requires `openpyxl` (pip install openpyxl).") from e
    wb = Workbook()
    ws_summary = wb.active
    ws_summary.title = "Summary"
    ws_summary.append(["Workspace", snap["workspace"]["name"]])
    ws_summary.append(["Target", snap["workspace"]["target"] or ""])
    ws_summary.append(["Generated", snap["report"]["generated_at"]])
    ws_summary.append(["Author", snap["report"]["author"]])
    ws_summary.append([])
    ws_summary.append(["Scans", snap["totals"]["scans"]])
    ws_summary.append(["Findings", snap["totals"]["findings"]])
    ws_summary.append(["Evidence", snap["totals"]["evidence"]])

    ws_findings = wb.create_sheet("Findings")
    ws_findings.append(["scan_tool", "category", "target", "severity", "title",
                        "host", "port", "service", "status", "cvss", "cve", "mitre"])
    for sc in snap['scans']:
        for f in sc['findings']:
            ws_findings.append([
                sc['tool'], sc['category'], sc['target'], f['severity'], f['title'],
                f['host'], f['port'], f['service'], f.get('status', 'open'),
                f.get('cvss_score'), f.get('cve_ids'), f.get('mitre_technique'),
            ])
    ws_ev = wb.create_sheet("Evidence")
    ws_ev.append(["filename", "sha256", "size_bytes", "mime_type", "created_at"])
    for ev in snap['evidence']:
        ws_ev.append([ev['filename'], ev['sha256'], ev['size_bytes'], ev['mime_type'], ev['created_at']])
    wb.save(str(out_path))


def _render_pdf(snap: dict, out_path: Path) -> None:
    """PDF generation. Tries weasyprint first (real PDF), falls back to a
    Markdown→HTML→PDF via reportlab as a plain-text PDF."""
    try:
        from weasyprint import HTML  # type: ignore
        HTML(string=_render_html(snap)).write_pdf(str(out_path))
        return
    except ImportError:
        pass
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
    except ImportError as e:
        raise RuntimeError(
            "PDF export requires `weasyprint` (best) or `reportlab` (fallback). "
            "pip install weasyprint  — or — pip install reportlab"
        ) from e
    c = canvas.Canvas(str(out_path), pagesize=LETTER)
    width, height = LETTER
    y = height - 50
    for line in _render_md(snap).splitlines():
        if y < 50:
            c.showPage(); y = height - 50
        c.drawString(40, y, line[:110])
        y -= 12
    c.save()
